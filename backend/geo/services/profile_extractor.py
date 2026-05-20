"""LLM 抽取品牌资料 — 把用户拖入的原始文本(简介 / 官网文案 / PRD 等)
结构化映射到 BrandProfile 字段。

复用与 content_generator 相同的 DeepSeek 双通道(直连 + OpenRouter fallback);
两个 key 都没配时调用方应给前端 503,而不是无声生成空资料。

入口:
- `extract_profile_from_text(text)` → (BrandProfile, model_id)
- `file_bytes_to_text(filename, data)` → str(用于 PDF / DOCX / 纯文本上传)

model_id == "" 表示 provider 不可用,调用方据此返回 503
"""
from __future__ import annotations

import io
import json
import logging
import os
import re

import requests

from geo.models.ai_telemetry import BrandProfile

log = logging.getLogger(__name__)

DEEPSEEK_BASE_URL = os.environ.get("DEEPSEEK_BASE_URL", "https://api.deepseek.com").rstrip("/")
DEEPSEEK_MODEL = os.environ.get("DEEPSEEK_MODEL", "deepseek-chat")
OPENROUTER_DEEPSEEK_MODEL = os.environ.get("OPENROUTER_DEEPSEEK_MODEL", "deepseek/deepseek-chat")
OPENROUTER_URL = "https://openrouter.ai/api/v1/chat/completions"
DEFAULT_TIMEOUT = int(os.environ.get("GEO_PROFILE_EXTRACT_TIMEOUT", "120"))


_SCHEMA_HINT = """\
请严格按下面 schema 输出 JSON(字段名 snake_case,缺失字段用空字符串 / 空数组,
不要凭空捏造):
{
  "profile_name": "...",
  "company_full_name": "...",
  "company_short_name": "...",
  "industry": "...",
  "core_business_lines": ["..."],
  "service_geo": "...",
  "creation_directions": ["..."],
  "copywriting_types": ["..."],
  "target_platforms": ["..."],
  "content_tones": ["..."],
  "content_redlines": ["..."],
  "team_size": "...",
  "founded_year": "...",
  "core_credentials": ["..."],
  "brand_diff_tags": ["..."],
  "core_service_overview": "...",
  "service_features": ["..."],
  "service_process": ["..."],
  "target_scenarios": ["..."],
  "service_guarantees": ["..."],
  "brand_story": "...",
  "key_person_story": "...",
  "case_stories": ["..."],
  "brand_values": "...",
  "available_materials": ["..."],
  "brand_slogan": "...",
  "core_message": "...",
  "extra_notes": "...",
  "seed_suggestions": ["...", "..."]
}

【核心原则:原文照搬,按段落分项,不删不缩】
所有 list 字段(包括但不限于 core_business_lines 核心业务线 / core_credentials 核心荣誉 /
case_stories 典型案例 / service_features / service_process / target_scenarios 等),都要遵守:
1. 原文里每一条独立的段落 / 项目符号 / 编号条目 = list 里独立的一项,一一对应,
   不要合并(比如原文 5 条荣誉,输出就必须是 5 条,不能压成 3 条)。
2. 每一项原文怎么写就怎么填,完整保留原句,不要缩写、不要总结、不要改写措辞;
   括号里的补充说明、年份、颁发机构、案例数据也都要带上。
3. 不要因为字数多就裁掉描述 — 哪怕一条荣誉占两三行字也要完整保留;
   除非超过 1000 字一条的极端长描述才适度截断。
4. 不要凭空合并相似条目;原文如果分成两段写,就分两条。
5. 反面例子:原文写「证券虚假陈述责任纠纷案件代理」,**不要**改写成
   「证券虚假陈述应对」/「证券虚假陈述」/「证券诉讼」等任何精简版本 —— 字字照搬原句。
   原文写「企业合规体系建设与刑事风险防控」,**不要**拆成
   「企业合规」+「刑事风险」两项,也**不要**简化成「合规建设」。

seed_suggestions 是给后续监测扩展用的「种子提示词」,要求:
- 3-8 条,每条 4-12 字
- 用真实用户会去问 AI 的口吻(短问 / 名词短语),不要写成长句
- 围绕本品牌的核心业务 / 用户痛点 / 行业关键词;避开品牌名本身
- 例子:跨境并购 / 律所推荐 / 遗产继承 / 跨境支付

只输出一个 JSON 对象,不要加 ``` 围栏,不要写解释文本。
"""

SYSTEM_PROMPT = (
    "你是品牌资料分析师,把用户给的原始资料(公司简介 / 官网文案 / PRD / 产品说明 / "
    "市场材料等)按【原文照搬、按段落分项】的方式映射到「品牌资料」JSON,"
    "并顺手给出几条可用作 AI 监测的「种子提示词」。\n"
    "你的任务是把原文里已有的内容搬到对应字段,而不是当编辑去改稿、压缩或润色 — "
    "用户后续会看到 list 里每一项,所以条目数和原文措辞必须保留。\n"
    + _SCHEMA_HINT
)


def _resolve_provider() -> tuple[str | None, str, str]:
    """returns (provider, model_id, api_key); provider=None 表示不可用."""
    ds_key = os.environ.get("DEEPSEEK_API_KEY", "").strip()
    or_key = os.environ.get("OPENROUTER_API_KEY", "").strip()
    if ds_key:
        return "deepseek", DEEPSEEK_MODEL, ds_key
    if or_key:
        return "openrouter", OPENROUTER_DEEPSEEK_MODEL, or_key
    return None, "", ""


def extract_profile_from_text(text: str) -> tuple[BrandProfile, list[str], str]:
    """returns (profile, seed_suggestions, model_id);model_id="" 表示 provider 不可用."""
    provider, model_id, api_key = _resolve_provider()
    if not provider:
        return BrandProfile(), [], ""

    if provider == "deepseek":
        url = f"{DEEPSEEK_BASE_URL}/chat/completions"
        headers = {
            "Authorization": f"Bearer {api_key}",
            "Content-Type": "application/json",
        }
    else:
        url = OPENROUTER_URL
        headers = {
            "Authorization": f"Bearer {api_key}",
            "Content-Type": "application/json",
            "HTTP-Referer": "https://www.vigilath.com",
            "X-Title": "GEO Profile Extractor",
        }
    payload = {
        "model": model_id,
        "messages": [
            {"role": "system", "content": SYSTEM_PROMPT},
            {"role": "user", "content": f"以下是原始资料,请抽取为品牌资料 JSON:\n\n{text}"},
        ],
        "temperature": 0.2,
        "response_format": {"type": "json_object"},
    }
    r = requests.post(url, json=payload, headers=headers, timeout=DEFAULT_TIMEOUT)
    if r.status_code >= 400:
        raise RuntimeError(f"{provider} {r.status_code}: {r.text[:200]}")
    data = r.json()
    try:
        content = data["choices"][0]["message"]["content"]
    except Exception as e:  # noqa: BLE001
        raise RuntimeError(f"unexpected {provider} response shape: {e}")
    parsed = _parse_json_loose(content)
    safe = _coerce_to_profile_fields(parsed)
    try:
        profile = BrandProfile(**safe)
    except Exception as e:  # noqa: BLE001
        # 真挂了再降级 — 之前所有 list 超长都会触发 max_length 校验失败,
        # 导致整张画像都空;coerce 之后这条路径基本不会进了
        log.warning("BrandProfile 映射失败,降级用空资料: %s; raw=%s", e, content[:300])
        profile = BrandProfile()
    # seed_suggestions:LLM 给的种子词候选,清洗 + 去重 + 限长(1-30 字,最多 8 条)
    raw_seeds = parsed.get("seed_suggestions") or []
    seeds: list[str] = []
    seen: set[str] = set()
    if isinstance(raw_seeds, list):
        for s in raw_seeds:
            if not isinstance(s, str):
                continue
            v = s.strip()
            if 1 <= len(v) <= 30 and v not in seen:
                seen.add(v)
                seeds.append(v)
            if len(seeds) >= 8:
                break
    return profile, seeds, model_id


def _coerce_to_profile_fields(parsed: dict) -> dict:
    """把 LLM 输出按 BrandProfile schema 强制兜底:
      - list 字段超 max_length 的截断到上限
      - LLM 给了 string 但 schema 要 list → 拆成单元素 list
      - LLM 给了 list 但 schema 要 string → join 成多行
      - None / 异常类型 → 丢弃(走默认值)
    防止单条字段 hard-fail 让整张画像归零.
    """
    out: dict = {}
    for name, field in BrandProfile.model_fields.items():
        if name not in parsed:
            continue
        raw = parsed[name]
        if raw is None:
            continue
        ann = field.annotation
        # 拿 max_length(Field(..., max_length=N))
        max_len = None
        for m in (field.metadata or []):
            ml = getattr(m, "max_length", None)
            if ml is not None:
                max_len = ml
                break

        is_list_field = _is_list_annotation(ann)
        if is_list_field:
            if isinstance(raw, list):
                items = [str(x).strip() for x in raw if x is not None and str(x).strip()]
            elif isinstance(raw, str):
                s = raw.strip()
                items = [s] if s else []
            else:
                continue
            if max_len is not None and len(items) > max_len:
                items = items[:max_len]
            out[name] = items
        else:
            # 字符串字段
            if isinstance(raw, list):
                s = "\n".join(str(x).strip() for x in raw if x is not None and str(x).strip())
            else:
                s = str(raw)
            if max_len is not None and len(s) > max_len:
                s = s[:max_len]
            out[name] = s
    return out


def _is_list_annotation(ann) -> bool:
    """判断 Pydantic 字段类型是否是 list[...]."""
    import typing
    origin = typing.get_origin(ann)
    return origin in (list, typing.List)


# ─────────────── 上传文件 → 纯文本 ─────────────────────────


# 上限同 BrandProfile.extra_notes 的合理边界:LLM 一次能消化的中文 ~30K 字
MAX_EXTRACTED_TEXT = 60_000

# 友好的错误消息体,调用方会包成 HTTPException
class FileParseError(Exception):
    pass


SUPPORTED_EXTS = ("txt", "md", "docx", "pdf")


def file_bytes_to_text(filename: str, data: bytes) -> str:
    """根据后缀分发到对应解析器;返回截断到 MAX_EXTRACTED_TEXT 的文本。

    支持:
    - .docx → python-docx 段落 + 表格
    - .pdf  → pypdf 按页 + 段落抽文字(扫描件 / 图片版 PDF 抽不到字,会抛 FileParseError)
    - .txt / .md → utf-8 / gbk 兜底

    其它类型(.doc / .csv / .json …)统一 415,提示用户复制内容到文本框,
    或另存为支持的格式。
    """
    name = (filename or "").lower()
    ext = name.rsplit(".", 1)[-1] if "." in name else ""

    if ext == "docx":
        return _truncate(_parse_docx(data))
    if ext == "pdf":
        return _truncate(_parse_pdf(data))
    if ext in ("txt", "md"):
        try:
            return _truncate(data.decode("utf-8"))
        except UnicodeDecodeError:
            try:
                return _truncate(data.decode("gbk", errors="replace"))
            except Exception as e:  # noqa: BLE001
                raise FileParseError(f"无法以 UTF-8 或 GBK 解码文件:{e}") from e
    if ext == "doc":
        raise FileParseError(
            "暂不支持 .doc 旧二进制格式,请用 Word 「另存为」.docx 后重试"
        )
    raise FileParseError(
        f"暂不支持的文件类型(.{ext or '未知'}),支持 {', '.join('.' + e for e in SUPPORTED_EXTS)}"
    )


def _truncate(text: str) -> str:
    text = (text or "").strip()
    if len(text) > MAX_EXTRACTED_TEXT:
        return text[:MAX_EXTRACTED_TEXT]
    return text


def _parse_docx(data: bytes) -> str:
    try:
        import docx  # python-docx
    except ImportError as e:
        raise FileParseError("后端缺少 python-docx 依赖,无法解析 Word") from e
    try:
        doc = docx.Document(io.BytesIO(data))
    except Exception as e:  # noqa: BLE001
        raise FileParseError(f"Word 解析失败(文件可能损坏):{e}") from e
    paras = [p.text for p in doc.paragraphs if p.text and p.text.strip()]
    # 表格也抽一下 — PRD / 公司简介常用表格列字段
    tables: list[str] = []
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text and c.text.strip()]
            if cells:
                tables.append(" | ".join(cells))
    out = "\n".join(paras + tables).strip()
    if not out:
        raise FileParseError("这份 Word 文档为空 / 没有可抽取的文字")
    return out


def _parse_pdf(data: bytes) -> str:
    """pypdf 按页抽文字,保留页内的段落换行 — LLM 后续靠这些换行识别"按段落分项"。

    扫描件 / 图片版 PDF 抽不到字符,会落到 FileParseError 让前端提示用户用 Word/文本。
    """
    try:
        from pypdf import PdfReader
    except ImportError as e:
        raise FileParseError("后端缺少 pypdf 依赖,无法解析 PDF") from e
    try:
        reader = PdfReader(io.BytesIO(data))
    except Exception as e:  # noqa: BLE001
        raise FileParseError(f"PDF 解析失败(文件可能损坏 / 加密):{e}") from e
    if getattr(reader, "is_encrypted", False):
        try:
            reader.decrypt("")
        except Exception:  # noqa: BLE001
            raise FileParseError("PDF 已加密,请先解除密码保护再上传")
    pages: list[str] = []
    for page in reader.pages:
        try:
            t = page.extract_text() or ""
        except Exception:  # noqa: BLE001
            t = ""
        t = t.strip()
        if t:
            pages.append(t)
    out = "\n\n".join(pages).strip()
    if not out:
        raise FileParseError(
            "这份 PDF 抽不到任何文字(可能是扫描件 / 图片版),"
            "请用 Word 重新导出 .docx,或复制内容粘贴到文本框"
        )
    return out


def _parse_json_loose(text: str) -> dict:
    s = (text or "").strip()
    if s.startswith("```"):
        s = s.lstrip("`").lstrip()
        if s.lower().startswith("json"):
            s = s[4:].lstrip()
        if s.endswith("```"):
            s = s[:-3].rstrip()
    try:
        v = json.loads(s)
        if isinstance(v, dict):
            return v
    except Exception:  # noqa: BLE001
        pass
    m = re.search(r"\{[\s\S]*\}", s)
    if m:
        try:
            v = json.loads(m.group(0))
            if isinstance(v, dict):
                return v
        except Exception:  # noqa: BLE001
            pass
    raise RuntimeError(f"failed to parse JSON from LLM output: {s[:200]}")
